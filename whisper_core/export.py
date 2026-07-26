"""Експорт розшифровки у субтитри (.srt/.vtt) і документ Word (.docx).

Вхід усюди — segments: список сегментів із таймкодами. Приймаємо три форми
(усі, що можуть прийти від рушія чи фейків): кортеж/список (start, end, text),
dict {"start","end","text"} або об'єкт із атрибутами start/end/text.

Норми субтитрів (за довідником, розділ 1.3 — BBC-консервативний профіль для
«бабусиної» аудиторії): 42 символи/рядок, максимум 2 рядки, ціль 15 CPS,
мінімум показу 0.83 с, максимум 7 с, паузи > 3 с — межа нового титру. Сирі
сегменти faster-whisper мерджимо (короткі суміжні), розриваємо рядок по
пунктуації/пробілу. SRT — UTF-8 БЕЗ BOM (BOM ламає старі DirectShow-фільтри),
кома в мілісекундах; VTT — заголовок WEBVTT, крапка в мілісекундах.
"""
from __future__ import annotations

import re
from pathlib import Path

MAX_LINE = 42          # символів у рядку субтитра
MAX_LINES = 2          # рядків у титрі
TARGET_CPS = 15.0      # знаків/с — комфортна швидкість читання
MIN_DUR = 0.83         # мінімальна тривалість показу титру, с
MAX_MERGE_GAP = 3.0    # пауза > неї — завжди межа нового титру, с
MAX_DUR = 7.0          # довше не тримаємо один титр, с


def _normalize(segments) -> list[dict]:
    """Будь-яка з трьох форм сегмента → [{"start","end","text"}] з float-часом.
    Порожні за текстом сегменти відкидаємо."""
    out = []
    for s in segments or []:
        if isinstance(s, dict):
            start, end, text = s.get("start"), s.get("end"), s.get("text", "")
        elif isinstance(s, (tuple, list)):
            start, end, text = s[0], s[1], s[2]
        else:
            start, end, text = s.start, s.end, s.text
        text = (text or "").strip()
        # без тексту або без таймкоду сегмент пропускаємо (float(None) інакше
        # зірвав би ВЕСЬ експорт TypeError'ом на одному кривому сегменті)
        if not text or start is None or end is None:
            continue
        out.append({"start": float(start), "end": float(end), "text": text})
    return out


def _fmt_ts(seconds: float, sep: str) -> str:
    """Секунди → HH:MM:SS{sep}mmm. sep = "," для SRT, "." для VTT.
    Округлення до мілісекунди робимо на цілих мс, щоб не було від'ємних
    мілісекунд чи «60» у полях (пастка ручної divmod-конвертації з довідника)."""
    if seconds < 0:
        seconds = 0.0
    total_ms = int(round(seconds * 1000))
    ms = total_ms % 1000
    total_s = total_ms // 1000
    s = total_s % 60
    m = (total_s // 60) % 60
    h = total_s // 3600
    return f"{h:02d}:{m:02d}:{s:02d}{sep}{ms:03d}"


def _wrap_two_lines(text: str, width: int = MAX_LINE) -> list[str]:
    """Розбити текст на рядки <= width по межах слів. Ціль — <= 2 рядки, але
    текст ніколи не втрачаємо: якщо не влазить — буде більше рядків (краще
    зайвий рядок, ніж зникле слово для немолодого читача)."""
    words = text.split()
    if not words:
        return [text]
    lines, cur = [], ""
    for w in words:
        cand = f"{cur} {w}".strip()
        if len(cand) <= width or not cur:
            cur = cand
        else:
            lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    # злити у 2 рядки якомога рівніше, якщо вийшло рівно 2 «природних» — лишаємо
    return lines


def _build_cues(segments) -> list[dict]:
    """Сирі сегменти → готові титри: мердж коротких, підтяг тривалості під CPS
    і мінімум, розбивка тексту на рядки. Повертає [{"start","end","lines"}]."""
    segs = _normalize(segments)
    if not segs:
        return []

    # 1) мердж: накопичуємо, доки титр закороткий і суміжний сегмент близько,
    #    і поки текст влазить у 2 рядки. Пауза > MAX_MERGE_GAP — завжди розрив.
    merged: list[dict] = []
    for seg in segs:
        if merged:
            prev = merged[-1]
            gap = seg["start"] - prev["end"]
            joined = f"{prev['text']} {seg['text']}"
            fits = len(_wrap_two_lines(joined)) <= MAX_LINES
            too_short = (prev["end"] - prev["start"]) < MIN_DUR
            within = (seg["end"] - prev["start"]) <= MAX_DUR
            if gap <= MAX_MERGE_GAP and fits and too_short and within:
                prev["text"] = joined
                prev["end"] = seg["end"]
                continue
        merged.append(dict(seg))

    # 2) тривалість: підтягнути під мінімум і під 15 CPS, не залазячи на
    #    наступний титр (лишаємо 40 мс проміжку). Це «подовжити показ до
    #    наступного сегмента» з довідника.
    cues = []
    for i, seg in enumerate(merged):
        start, end, text = seg["start"], seg["end"], seg["text"]
        need_cps = start + len(text) / TARGET_CPS      # скільки треба під 15 CPS
        want_end = max(end, start + MIN_DUR, need_cps)
        nxt = merged[i + 1]["start"] if i + 1 < len(merged) else None
        if nxt is not None:
            want_end = min(want_end, nxt - 0.04)
        end = max(end, min(want_end, start + MAX_DUR))
        if end <= start:                                # захист від виродженого титру
            end = start + MIN_DUR
        cues.append({"start": start, "end": end, "lines": _wrap_two_lines(text)})
    return cues


def to_srt(segments) -> str:
    """Segments → рядок SRT (UTF-8 без BOM пише вже викликач). Кома в мс,
    години завжди, порожній рядок між блоками і в кінці."""
    cues = _build_cues(segments)
    blocks = []
    for i, c in enumerate(cues, 1):
        ts = f"{_fmt_ts(c['start'], ',')} --> {_fmt_ts(c['end'], ',')}"
        blocks.append(f"{i}\n{ts}\n" + "\n".join(c["lines"]))
    return "\n\n".join(blocks) + "\n"


def to_vtt(segments) -> str:
    """Segments → рядок WebVTT. Заголовок WEBVTT + порожній рядок, крапка в мс."""
    cues = _build_cues(segments)
    parts = ["WEBVTT", ""]
    for c in cues:
        ts = f"{_fmt_ts(c['start'], '.')} --> {_fmt_ts(c['end'], '.')}"
        parts.append("")
        parts.append(ts)
        parts.extend(c["lines"])
    return "\n".join(parts) + "\n"


# --- feature/edit-pack: імпорт субтитрів назад (SRT/VTT → сегменти) ---

# Рядок таймкоду: «ГГ:ХХ:СС(,|.)ммм --> ...». Години в SRT/VTT опційні (VTT
# дозволяє ХХ:СС.ммм), розділювач мс — кома (SRT) або крапка (VTT) — приймаємо
# обидва в будь-якому форматі, бо ручні редактори субтитрів їх плутають.
_TS = r"(?:(\d+):)?(\d{1,2}):(\d{2})[.,](\d{1,3})"
_CUE_RE = re.compile(rf"^\s*{_TS}\s*-->\s*{_TS}")


def _ts_to_seconds(h, m, s, ms) -> float:
    """Групи регексу таймкоду → секунди (float). Години можуть бути None; частка
    секунди трактується як десятковий дріб праворуч: «.5»→500 мс, «.05»→50 мс."""
    ms = (ms + "000")[:3]
    return (int(h or 0) * 3600 + int(m) * 60 + int(s)) + int(ms) / 1000.0


def parse_subtitles(content: str) -> list[dict]:
    """Текст .srt або .vtt → [{"start","end","text"}] у порядку появи.

    Один парсер на обидва формати (вони відрізняються лише заголовком WEBVTT,
    розділювачем мс і номерами блоків, а структура «таймкод + рядки тексту»
    спільна). Стійкий до країв: BOM на початку, заголовок WEBVTT, номер блоку
    перед таймкодом (SRT), налаштування VTT-кʼю після таймкоду, блоки NOTE,
    багаторядковий текст (склеюємо через пробіл), порожні блоки (пропускаємо),
    CRLF/CR. Порожній за текстом кʼю відкидаємо. Час завжди в секундах-float."""
    if not content:
        return []
    content = content.lstrip("﻿")                    # BOM
    lines = content.replace("\r\n", "\n").replace("\r", "\n").split("\n")

    cues: list[dict] = []
    i = 0
    n = len(lines)
    while i < n:
        m = _CUE_RE.match(lines[i])
        if not m:
            i += 1
            continue
        start = _ts_to_seconds(m.group(1), m.group(2), m.group(3), m.group(4))
        end = _ts_to_seconds(m.group(5), m.group(6), m.group(7), m.group(8))
        i += 1
        text_lines = []
        # текст — до наступного порожнього рядка або наступного таймкоду
        while i < n and lines[i].strip() and not _CUE_RE.match(lines[i]):
            text_lines.append(lines[i].strip())
            i += 1
        text = " ".join(text_lines).strip()
        if text:
            cues.append({"start": start, "end": end, "text": text})
    return cues


# --- feature/auto-export: дозапис розшифровок у файл-день теки ---

def append_transcript(directory, text: str, when, fmt: str = "md"):
    """Дописати одну завершену розшифровку у файл-день теки й повернути його шлях.

    Файл — <тека>/balachky-РРРР-ММ-ДД.<fmt> (один на день, дозапис у кінець,
    UTF-8). md: заголовок «## ГГ:ХХ», текст, порожній рядок. txt: рядок
    «[ГГ:ХХ] текст» + порожній рядок. `when` — datetime (передається
    параметром заради тестовності: дає і дату для імені, і час для заголовка).
    Порожній текст ігноруємо (None). Тека має існувати — якщо її нема чи нема
    прав, open() підніме OSError, і викликач сам вирішить, як попередити."""
    text = (text or "").strip()
    if not text:
        return None
    if fmt not in ("md", "txt"):
        fmt = "md"
    day = when.strftime("%Y-%m-%d")
    hm = when.strftime("%H:%M")
    path = Path(directory) / f"balachky-{day}.{fmt}"
    if fmt == "txt":
        entry = f"[{hm}] {text}\n\n"
    else:
        entry = f"## {hm}\n{text}\n\n"
    # newline="" — без нього текстовий режим на Windows перетворив би кожен \n
    # на CRLF; ми ж пишемо готові рядки з чистим LF (як у докстрінгу й тестах).
    with open(path, "a", encoding="utf-8", newline="") as f:
        f.write(entry)
    return path


# --- feature/markdown-export: Markdown з YAML-frontmatter (для Obsidian) ---

def _yaml_quote(value) -> str:
    """Рядок → YAML double-quoted літерал. Екрануємо \\ і " (шлях/назва з ними
    інакше зламали б frontmatter). Загортаємо навіть «безпечні» значення на кшталт
    тривалості «00:03:12»: без лапок YAML 1.1 прочитав би її як число за основою 60."""
    return '"' + str(value).replace("\\", "\\\\").replace('"', '\\"') + '"'


def duration_str(seconds) -> "str | None":
    """Секунди → «ГГ:ХХ:СС» (для поля duration у frontmatter). None → None."""
    if seconds is None:
        return None
    total = int(round(float(seconds)))
    if total < 0:
        total = 0
    return f"{total // 3600:02d}:{(total // 60) % 60:02d}:{total % 60:02d}"


def build_frontmatter(meta) -> str:
    """meta → YAML-блок між «---». Ключі у фіксованому порядку; порожні
    пропускаємо. meta: {"date","type","source","duration","tags"}. tags — типово
    [балачки]; type (напр. "meeting" для каналу Obsidian) — просте скалярне поле,
    корисне для запитів у сховищі; source і duration беремо в лапки (_yaml_quote)."""
    meta = meta or {}
    lines = ["---"]
    date = meta.get("date")
    if date:
        lines.append(f"date: {date}")
    note_type = meta.get("type")
    if note_type:
        lines.append(f"type: {note_type}")
    source = meta.get("source")
    if source:
        lines.append(f"source: {_yaml_quote(source)}")
    duration = meta.get("duration")
    if duration:
        lines.append(f"duration: {_yaml_quote(duration)}")
    tags = meta.get("tags") or ["балачки"]
    lines.append("tags: [" + ", ".join(tags) + "]")
    lines.append("---")
    return "\n".join(lines)


def _segments_duration(segments) -> "str | None":
    """Тривалість розшифровки з таймкодів сегментів (кінець останнього) або None,
    якщо сегментів без таймкоду немає (напр. коротка розшифровка з диктування)."""
    segs = _normalize(segments)
    if not segs:
        return None
    return duration_str(max(s["end"] for s in segs))


def to_markdown(text: str, meta: dict, segments=None) -> str:
    """Розшифровка → Markdown: YAML-frontmatter + суцільний текст. duration
    підставляємо з сегментів, якщо є таймкоди й meta його не задав. Текст ніколи
    не втрачаємо — навіть без сегментів пишемо тіло як є."""
    m = dict(meta or {})
    if not m.get("duration"):
        d = _segments_duration(segments)
        if d:
            m["duration"] = d
    return build_frontmatter(m) + "\n\n" + (text or "").strip() + "\n"


def to_docx(segments, meta: dict, path, fallback_text: str = "") -> None:
    """Segments → .docx зі статутною розміткою (statut_docx): заголовок (ім'я
    файлу), дата, далі абзаци «[гг:хх:сс] текст». meta: {"filename": str,
    "date": str}. Форматування (A4, береги, Times New Roman 14, одинарний,
    відступ 1,25 см) — дефолт статутного документа. Без сегментів (напр. коротка
    розшифровка без часових позначок) пише fallback_text одним абзацом — документ
    ніколи не лишається порожнім із хибним «Збережено»."""
    from whisper_core import statut_docx as st

    doc = st.new_document()
    st.add_heading(doc, meta.get("filename", "Розшифровка"), level=1)
    date_str = meta.get("date", "")
    if date_str:
        st.add_body(doc, f"Дата: {date_str}")

    norm = _normalize(segments)
    if norm:
        for seg in norm:
            p = doc.add_paragraph()          # успадковує Normal (статутний стиль)
            stamp = p.add_run(f"[{_fmt_ts(seg['start'], '.').split('.')[0]}] ")
            stamp.bold = True
            p.add_run(st.sanitize(seg["text"]))
    elif fallback_text.strip():
        st.add_body(doc, fallback_text.strip())

    doc.save(str(path))


def protocol_to_docx(markdown: str, path, *, title: str = "Протокол наради") -> None:
    """Markdown-протокол наради (feature/ai-protocol) → .docx. Розбирає прості
    елементи, що їх дає LLM за нашим форматом:
      «## …» → заголовок 2 рівня;
      таблиця «| a | b |» (з рядком-розділювачем «|---|») → docx-таблиця;
      «- …» → пункт із літеральним маркером «– » (авто-numPr Word не друкує!);
    інше — абзац. Складніший Markdown не інтерпретуємо — рядок іде як звичайний
    текст (документ ніколи не порожній/битий). Форматування — статутне
    (statut_docx): A4, береги, Times New Roman 14, одинарний, відступ 1,25 см."""
    from whisper_core import statut_docx as st

    doc = st.new_document()
    st.add_heading(doc, title, level=1)

    lines = (markdown or "").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("## "):
            st.add_heading(doc, stripped[3:].strip(), level=2)
            i += 1
        elif stripped.startswith("|") and stripped.endswith("|"):
            # зібрати весь блок таблиці (суміжні рядки з «|»)
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i].strip())
                i += 1
            _md_table_to_docx(doc, block)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            # маркер — літеральний «– », НЕ style="List Bullet" (numPr не друкується)
            st.add_body(doc, "– " + stripped[2:].strip())
            i += 1
        else:
            st.add_body(doc, stripped)
            i += 1

    doc.save(str(path))


def _md_table_to_docx(doc, block) -> None:
    """Рядки Markdown-таблиці → docx-таблиця. Рядок-розділювач (|---|---|)
    пропускаємо. Порожній/однорядковий блок → нічого (без краху)."""
    def cells(row):
        return [c.strip() for c in row.strip().strip("|").split("|")]

    rows = [cells(r) for r in block]
    # відкинути рядок-розділювач «|---|---|» і порожні рядки
    rows = [r for r in rows
            if any(c for c in r) and not all(set(c) <= {"-", ":", " "} for c in r)]
    if not rows:
        return
    from docx.shared import Pt

    from whisper_core import statut_docx as st

    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncols)
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        pass
    for r in rows:
        row_cells = table.add_row().cells
        for j in range(ncols):
            cell = row_cells[j]
            cell.text = st.sanitize(r[j]) if j < len(r) else ""
            # стиль таблиці може задавати власний шрифт — прибиваємо TNR 14 явно
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = st.FONT_NAME
                    run.font.size = Pt(st.FONT_SIZE_PT)
