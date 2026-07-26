"""Статутна розмітка Word-документів («Службовий документ»).

Дефолтне форматування docx-експортів під українське військове діловодство —
щоб вихідний файл одразу виглядав правильно, без ручного переоформлення:

- аркуш A4, книжна орієнтація;
- береги: лівий 30 мм, правий 10 мм, верхній 20 мм, нижній 20 мм;
- шрифт Times New Roman 14 pt скрізь;
- міжрядковий одинарний (line=240, lineRule=auto), інтервал до/після абзацу = 0;
- абзацний відступ першого рядка 1,25 см (firstLine=709 twips);
- текст по ширині.

Два правила тексту документа:
- нумерація/маркери пунктів — ЛИШЕ літеральним текстом («1.», «– »); авто-нумерація
  Word (w:numPr) НЕ ДРУКУЄТЬСЯ на принтері, тож її тут немає взагалі;
- у тексті прямі лапки замість «»/“”, тире коротке «–» (en-dash), не довге «—» (em-dash).

Модуль нейтральний: жодних згадок нормативних актів у самому документі — це просто
типове форматування. Використовується як спільна основа всіх docx-експортів.
"""
from __future__ import annotations

FONT_NAME = "Times New Roman"
FONT_SIZE_PT = 14

# «»‹›„“”‟ — усі кутові/криві лапки → пряма ". Em-dash та довші тире → en-dash «–».
_QUOTES = "«»‹›„“”‟"
_QUOTES_MAP = {ord(c): '"' for c in _QUOTES}
_DASHES_MAP = {ord("—"): "–", ord("―"): "–"}  # em-dash, horizontal bar → en-dash


def sanitize(text: str) -> str:
    """Прибрати з тексту документа заборонені за статутним оформленням символи:
    кутові/криві лапки → прямі ", довге тире «—» → коротке «–». None/порожнє →
    порожній рядок (щоб виклик у ланцюжку не падав)."""
    if not text:
        return ""
    return str(text).translate(_QUOTES_MAP).translate(_DASHES_MAP)


def new_document():
    """Порожній Document зі статутною розміткою: сторінка A4 з берегами й стиль
    Normal (Times New Roman 14, одинарний без відбивок, відступ 1,25 см, по ширині).
    Далі додавати абзаци через add_body/add_heading або звичайним doc.add_paragraph
    (успадкує Normal)."""
    from docx import Document

    doc = Document()
    _setup_page(doc)
    _setup_normal(doc)
    return doc


def _setup_page(doc) -> None:
    from docx.enum.section import WD_ORIENT
    from docx.shared import Mm

    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.left_margin = Mm(30)
    sec.right_margin = Mm(10)
    sec.top_margin = Mm(20)
    sec.bottom_margin = Mm(20)


def _setup_normal(doc) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    st = doc.styles["Normal"]
    st.font.name = FONT_NAME
    st.font.size = Pt(FONT_SIZE_PT)
    pf = st.paragraph_format
    # float → line=240, lineRule=auto (одинарний, як вимагає розмітка)
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(1.25)          # firstLine=709 twips
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_body(doc, text: str):
    """Абзац тексту документа: успадковує Normal (по ширині, відступ першого
    рядка), текст очищується sanitize()."""
    return doc.add_paragraph(sanitize(text))


def add_heading(doc, text: str, level: int = 1):
    """Заголовок статутного документа — жирний абзац Times New Roman 14 БЕЗ
    абзацного відступу (перший рівень — по центру, глибші — зліва). Свідомо НЕ
    вживаємо doc.add_heading: вбудовані стилі Heading дають інший шрифт/кегль/колір,
    що ламає «TNR 14 скрізь»."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if level <= 1
                   else WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(sanitize(text))
    run.bold = True
    return p
