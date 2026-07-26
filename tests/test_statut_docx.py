# feature/docx-statut-layout
"""Тести статутної розмітки docx-експортів (whisper_core.statut_docx) і її
застосування у export.to_docx / export.protocol_to_docx. Перевіряємо ЗГЕНЕРОВАНИЙ
файл — читанням через python-docx та розбором word/document.xml, — а не лише
виклики: береги, шрифт/кегль, міжрядковий, відступ першого рядка, вирівнювання,
відсутність авто-нумерації (w:numPr), очищення лапок «» і довгого тире «—».
"""
import os
import tempfile
import unittest
import zipfile

from docx import Document
from docx.shared import Mm, Pt, Twips

from whisper_core import export, statut_docx


def _document_xml(path: str) -> str:
    with zipfile.ZipFile(path) as z:
        return z.read("word/document.xml").decode("utf-8")


class SanitizeTests(unittest.TestCase):
    def test_guillemets_to_straight(self):
        self.assertEqual(statut_docx.sanitize("він «тут» і ‹там›"), 'він "тут" і "там"')

    def test_curly_quotes_to_straight(self):
        self.assertEqual(statut_docx.sanitize("„низ“ і “верх”"), '"низ" і "верх"')

    def test_em_dash_to_en_dash(self):
        # довге тире — та горизонтальна риска ― → коротке –
        self.assertEqual(statut_docx.sanitize("а — б ― в"), "а – б – в")

    def test_none_and_empty(self):
        self.assertEqual(statut_docx.sanitize(None), "")
        self.assertEqual(statut_docx.sanitize(""), "")


class NewDocumentLayoutTests(unittest.TestCase):
    def setUp(self):
        self.doc = statut_docx.new_document()

    def test_page_size_a4_portrait(self):
        # порівнюємо у twips: Mm→EMU→twips-округлення (у файлі — цілі twips)
        sec = self.doc.sections[0]
        self.assertEqual(sec.page_width.twips, Mm(210).twips)     # 11906
        self.assertEqual(sec.page_height.twips, Mm(297).twips)    # 16838

    def test_margins_30_10_20_20(self):
        sec = self.doc.sections[0]
        self.assertEqual(sec.left_margin.twips, Mm(30).twips)     # 1701
        self.assertEqual(sec.right_margin.twips, Mm(10).twips)    # 567
        self.assertEqual(sec.top_margin.twips, Mm(20).twips)      # 1134
        self.assertEqual(sec.bottom_margin.twips, Mm(20).twips)

    def test_normal_font_tnr_14(self):
        normal = self.doc.styles["Normal"]
        self.assertEqual(normal.font.name, "Times New Roman")
        self.assertEqual(normal.font.size, Pt(14))

    def test_normal_spacing_single_no_gaps(self):
        pf = self.doc.styles["Normal"].paragraph_format
        self.assertEqual(pf.line_spacing, 1.0)
        self.assertEqual(pf.space_before, Pt(0))
        self.assertEqual(pf.space_after, Pt(0))

    def test_normal_first_line_indent_709(self):
        pf = self.doc.styles["Normal"].paragraph_format
        self.assertEqual(pf.first_line_indent, Twips(709))

    def test_body_paragraph_inherits_and_sanitizes(self):
        statut_docx.add_body(self.doc, "текст «А» — Б")
        p = self.doc.paragraphs[-1]
        self.assertEqual(p.text, 'текст "А" – Б')

    def test_heading_no_indent_bold(self):
        statut_docx.add_heading(self.doc, "ЗАГОЛОВОК", level=1)
        p = self.doc.paragraphs[-1]
        self.assertEqual(p.paragraph_format.first_line_indent, Twips(0))
        self.assertTrue(p.runs[0].bold)


class ToDocxFileTests(unittest.TestCase):
    def setUp(self):
        self.dir = tempfile.mkdtemp()
        self.path = os.path.join(self.dir, "t.docx")

    def test_transcript_file_layout_and_sanitized(self):
        export.to_docx(
            [{"start": 0, "end": 2, "text": 'Він «привіт» — і пішов'}],
            {"filename": "Розмова", "date": "18.07.2026"}, self.path)
        xml = _document_xml(self.path)
        # немає авто-нумерації і заборонених символів
        self.assertNotIn("w:numPr", xml)
        self.assertNotIn("«", xml)
        self.assertNotIn("»", xml)
        self.assertNotIn("—", xml)
        # береги (twips) і A4
        self.assertIn('w:left="1701"', xml)
        self.assertIn('w:right="567"', xml)
        self.assertIn('w:top="1134"', xml)
        self.assertIn('w:bottom="1134"', xml)
        self.assertIn('w:w="11906"', xml)
        # текст сегмента дійшов очищеним
        doc = Document(self.path)
        full = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn('Він "привіт" – і пішов', full)

    def test_fallback_text_when_no_segments(self):
        export.to_docx([], {"filename": "Порожньо"}, self.path,
                       fallback_text="лише текст")
        doc = Document(self.path)
        self.assertIn("лише текст", "\n".join(p.text for p in doc.paragraphs))


class ProtocolDocxFileTests(unittest.TestCase):
    def setUp(self):
        self.dir = tempfile.mkdtemp()
        self.path = os.path.join(self.dir, "p.docx")
        md = ("## Розділ\n"
              "- пункт «А» — раз\n"
              "- пункт Б\n\n"
              "| Час | Подія |\n|---|---|\n| 10:00 | старт «X» |\n\n"
              "Звичайний абзац.")
        export.protocol_to_docx(md, self.path, title="Службовий документ")

    def test_no_numpr_even_with_bullets(self):
        # ключове: маркери «-» стали літеральним «– », без style="List Bullet"/numPr
        self.assertNotIn("w:numPr", _document_xml(self.path))

    def test_bullets_rendered_as_literal_dash(self):
        doc = Document(self.path)
        full = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("– пункт", full)      # літеральний en-dash-маркер

    def test_sanitized_and_layout(self):
        xml = _document_xml(self.path)
        self.assertNotIn("«", xml)
        self.assertNotIn("—", xml)
        self.assertIn('w:left="1701"', xml)

    def test_table_cells_are_tnr14(self):
        doc = Document(self.path)
        self.assertTrue(doc.tables, "таблиця має бути в документі")
        run = doc.tables[0].rows[0].cells[0].paragraphs[0].runs[0]
        self.assertEqual(run.font.name, "Times New Roman")
        self.assertEqual(run.font.size, Pt(14))


if __name__ == "__main__":
    unittest.main()
